# %%
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound
from pytube import YouTube
import re
import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer, pipeline
from docx import Document

# Fetch the transcript and video title
video_id = '8_DBT0XFQMQ'  # Replace with your YouTube video ID
video_url = f"https://www.youtube.com/watch?v={video_id}"

# Get video title
yt = YouTube(video_url)
video_title = yt.title

try:
    transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
except NoTranscriptFound:
    print("No Spanish transcript found.")
    transcript = []

# Handle the transcript data
transcript_text = ""
for entry in transcript:
    transcript_text += f"{entry['start']} - {entry.get('end', entry['start'] + 5)}: {entry['text']}\n"

# Clean the transcript
cleaned_lines = []
for line in transcript_text.split('\n'):
    cleaned_line = re.sub(r'^\d+\.\d+ - \d+\.\d+: ', '', line)
    cleaned_lines.append(cleaned_line)
cleaned_text = " ".join(cleaned_lines).strip()

# Summarize the cleaned transcript using the Hugging Face Transformers library (Spanish)
print(f"Versión de PyTorch: {torch.__version__}")
device = 0 if torch.cuda.is_available() else -1
print(f"Usando dispositivo: {'GPU' if device == 0 else 'CPU'}")

# Split the content into smaller chunks
max_chunk_size = 1000  # Adjust the chunk size as needed
chunks = [cleaned_text[i:i + max_chunk_size] for i in range(0, len(cleaned_text), max_chunk_size)]

# Initialize the model and tokenizer
model_name = "mrm8488/bert2bert_shared-spanish-finetuned-summarization"
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
tokenizer = AutoTokenizer.from_pretrained(model_name)

# Initialize the summarization pipeline
summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, device=device)

def split_text(text, max_length):
    """Splits text into chunks of max_length."""
    return [text[i:i+max_length] for i in range(0, len(text), max_length)]

# Generate the summary for each chunk and concatenate the results
summary_text = ""
for chunk in chunks:
    summary = summarizer(chunk, max_length=150, min_length=50, do_sample=False)
    summary_text += summary[0]['summary_text'] + " "

# Split the concatenated summary text into smaller chunks
max_model_length = 512  # Adjust based on the model's max input length
summary_chunks = split_text(summary_text, max_model_length)

# Generate a summary for each smaller chunk
final_summary_text = ""
for chunk in summary_chunks:
    final_summary = summarizer(chunk, max_length=150, min_length=50, do_sample=False)
    final_summary_text += final_summary[0]['summary_text'] + " "

# Create a Word document and save the cleaned transcript
transcript_doc = Document()
transcript_doc.add_heading(video_title, 0)
transcript_doc.add_heading('Transcripción Limpia', level=1)
transcript_doc.add_paragraph(cleaned_text)

transcript_output_file_path = f"{video_title}_transcript.docx"
transcript_doc.save(transcript_output_file_path)

# Create another Word document and save the summary
summary_doc = Document()
summary_doc.add_heading(video_title, 0)
summary_doc.add_heading('Resumen', level=1)
summary_doc.add_paragraph(final_summary_text.strip())

summary_output_file_path = f"{video_title}_summary.docx"
summary_doc.save(summary_output_file_path)

print(f"Transcripción limpia guardada en {transcript_output_file_path}")
print(f"Resumen guardado en {summary_output_file_path}")